"""
Report generation utilities:
- generate_excel_report: main multi-sheet Excel report (analysis, summary, PubMed)
- generate_word_report: compact human-readable Word document without URLs
"""

import pandas as pd
from pathlib import Path
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

log = logging.getLogger(__name__)

def generate_excel_report(all_verified_data: list[dict], output_path: Path, protocol_summary: str = None):
    """
    Generate an enhanced .xlsx report from the fully verified and synthesized data.
    Includes dosage, route, research links, evidence justification, and protocol summary.

    Args:
        all_verified_data: The list of verification records, each expected to
                           have the 'system_ud' key added.
        output_path: The Path object for the output .xlsx file.
        protocol_summary: Optional protocol executive summary.
    """
    if not all_verified_data:
        log.warning("No data provided to generate a report. Skipping file creation.")
        return

    log.info(f"Generating enhanced Excel report for {len(all_verified_data)} records...")

    try:
        # Prepare data for the DataFrame
        # Построение основной строки отчета по каждому препарату
        report_data = []
        for r in all_verified_data:
            # Enhanced status summary with detailed formulary info
            fda_data = r.get('approval_fda', {})
            is_fda_approved = fda_data.get('approved', False) if isinstance(fda_data, dict) else fda_data
            
            # Kazakhstan register info - detailed
            kz_data = r.get('kz_register', {})
            kz_status_str = ''
            kz_display = '❌ Не найден'
            if isinstance(kz_data, dict) and kz_data.get('found'):
                status_parts = []
                status_parts.append("✅ Зарегистрирован в РК")
                
                # Status from "Вид" column
                reg_status = kz_data.get('registration_status', '').strip()
                if reg_status:
                    status_parts.append(f"Статус: {reg_status}")
                
                # Reregistration status
                rereg_status = kz_data.get('reregistration_status', '').strip()
                if rereg_status:
                    status_parts.append(f"Перерегистрация: {rereg_status}")
                
                # Registration number
                reg_number = kz_data.get('registration_number', '').strip()
                if reg_number:
                    status_parts.append(f"Номер: {reg_number}")
                
                # Dosage form from "Лек. форма" column
                dosage_form = kz_data.get('dosage_forms', '').strip()
                if dosage_form:
                    status_parts.append(f"Форма: {dosage_form}")
                
                kz_status_str = '. '.join(status_parts) + '.'
                
                # Compact display for status summary with emoji
                kz_display_parts = ["✅ Зарегистрирован"]
                if reg_status:
                    kz_display_parts.append(reg_status)
                if dosage_form:
                    kz_display_parts.append(f"({dosage_form})")
                kz_display = ' '.join(kz_display_parts)
            
            # Enhanced status summary with emojis (no confidence percentage)
            status_parts_list = []
            
            # BNF
            if r.get('formulary_bnf'):
                status_parts_list.append("📚 BNF: ✅ Да")
            else:
                status_parts_list.append("📚 BNF: ❌ Нет")
            
            # BNF Children
            if r.get('formulary_bnfc'):
                status_parts_list.append("👶 BNF_Children: ✅ Да")
            else:
                status_parts_list.append("👶 BNF_Children: ❌ Нет")
            
            # WHO EML
            if r.get('formulary_who_eml'):
                status_parts_list.append("🌍 WHO EML: ✅ Да")
            else:
                status_parts_list.append("🌍 WHO EML: ❌ Нет")
            
            # FDA
            if is_fda_approved:
                status_parts_list.append("🇺🇸 FDA: ✅ Одобрено")
            else:
                status_parts_list.append("🇺🇸 FDA: ❌ Не одобрено")
            
            # EMA
            ema_status = r.get('approval_ema', '')
            if ema_status == 'Approved':
                status_parts_list.append("🇪🇺 EMA: ✅ Одобрено")
            elif ema_status:
                status_parts_list.append(f"🇪🇺 EMA: {ema_status}")
            else:
                status_parts_list.append("🇪🇺 EMA: ❌ Не найдено")
            
            # Kazakhstan
            status_parts_list.append(f"🇰🇿 Казахстан: {kz_display}")
            
            status_summary = ' | '.join(status_parts_list)

            # PubMed articles links
            pubmed_data = r.get('evidence_pubmed', {})
            articles = pubmed_data.get('articles', []) if isinstance(pubmed_data, dict) else []
            total_found = pubmed_data.get('total_found', 0) if isinstance(pubmed_data, dict) else len(articles) if articles else 0
            meta_count = pubmed_data.get('meta_analysis_count', 0) if isinstance(pubmed_data, dict) else 0
            systematic_count = pubmed_data.get('systematic_review_count', 0) if isinstance(pubmed_data, dict) else 0
            rct_count = pubmed_data.get('rct_count', 0) if isinstance(pubmed_data, dict) else 0
            
            pubmed_links = ''
            if articles:
                # Build links from articles with valid title and url
                valid_articles = [a for a in articles[:5] if a.get('title') and a.get('url')]
                if valid_articles:
                    links = [f"{a.get('title', '')[:50]}... - {a.get('url', '')}" for a in valid_articles]
                    pubmed_links = ' | '.join(links)
            
            # Build PubMed info string - show count even if no links
            if total_found > 0 or meta_count > 0 or systematic_count > 0 or rct_count > 0:
                pubmed_info_parts = []
                if total_found > 0:
                    pubmed_info_parts.append(f"Всего найдено: {total_found}")
                if meta_count > 0:
                    pubmed_info_parts.append(f"Мета-анализы: {meta_count}")
                if systematic_count > 0:
                    pubmed_info_parts.append(f"Систематические обзоры: {systematic_count}")
                if rct_count > 0:
                    pubmed_info_parts.append(f"РКИ: {rct_count}")
                
                pubmed_display = f"✅ {', '.join(pubmed_info_parts)}"
                if pubmed_links:
                    pubmed_display += f" | Ссылки: {pubmed_links}"
            else:
                pubmed_display = '❌ Не найдено'
            
            # Extract PubMed counts for summary generation (already extracted above)
            
            # Clinical trials links
            trials = r.get('clinical_trials', [])
            trials_count = len(trials) if trials else 0
            trials_links = ''
            if trials:
                # Build links from trials with valid title and url
                valid_trials = [t for t in trials[:3] if t.get('title') and t.get('url')]
                if valid_trials:
                    links = [f"{t.get('title', '')[:50]}... - {t.get('url', '')}" for t in valid_trials]
                    trials_links = ' | '.join(links)
            
            # Build ClinicalTrials info string - show count even if no links
            if trials_count > 0:
                trials_display = f"✅ Найдено исследований: {trials_count}"
                if trials_links:
                    trials_display += f" | Ссылки: {trials_links}"
            else:
                trials_display = '❌ Не найдено'

            # Comments with all issues
            comments_parts = []
            if r.get('comments'):
                comments_parts.append(r.get('comments'))
            if kz_status_str:
                comments_parts.append(kz_status_str)
            
            comments = '; '.join(comments_parts) if comments_parts else ''

            # System UD with justification and emoji
            system_ud = r.get('system_ud', 'D')
            if isinstance(system_ud, tuple):
                system_ud_level = system_ud[0]
            else:
                system_ud_level = str(system_ud)
            
            # Add emoji to GRADE level (4 levels: A, B, C, D)
            grade_emoji_map = {
                'A': '🟢',
                'B': '🟡',
                'C': '🟠',
                'D': '🔴'
            }
            grade_emoji = grade_emoji_map.get(system_ud_level, '⚪')
            system_ud_display = f"{grade_emoji} {system_ud_level}"
            
            justification = r.get('evidence_level_justification', '')

            # Summary recommendation (from Gemini analysis if available, otherwise from justification)
            summary_recommendation = r.get('summary_recommendation', '')
            if not summary_recommendation:
                # Fallback: create summary from available data with emojis
                summary_parts = []
                if system_ud_level and system_ud_level != 'D':
                    summary_parts.append(f"{grade_emoji} Уровень доказательности: {system_ud_level}")
                if is_fda_approved:
                    summary_parts.append("✅ Одобрен FDA")
                if r.get('formulary_bnf') or r.get('formulary_who_eml'):
                    summary_parts.append("📚 Включен в формуляры")
                # Always show research counts if available (use total_found from earlier)
                if total_found > 0:
                    research_info = []
                    if meta_count > 0:
                        research_info.append(f"мета-анализы: {meta_count}")
                    if systematic_count > 0:
                        research_info.append(f"систематические обзоры: {systematic_count}")
                    if rct_count > 0:
                        research_info.append(f"РКИ: {rct_count}")
                    if research_info:
                        summary_parts.append(f"📊 Найдены исследования ({', '.join(research_info)}, всего: {total_found})")
                    else:
                        summary_parts.append(f"📊 Найдено исследований: {total_found}")
                elif meta_count > 0 or systematic_count > 0:
                    # Fallback for old data format
                    summary_parts.append(f"📊 Найдены мета-анализы/систематические обзоры ({meta_count + systematic_count})")
                if summary_parts:
                    summary_recommendation = '. '.join(summary_parts) + '.'
                else:
                    summary_recommendation = justification if justification else '❓ Недостаточно данных для составления резюме.'

            # Extract additional fields from HTML demo version
            completeness = r.get('completeness', '')
            compliance_notes = r.get('compliance_notes', '')
            evidence_commentary = r.get('evidence_commentary', {})
            
            # Build additional info from evidence_commentary
            additional_info_parts = []
            if completeness:
                additional_info_parts.append(f"Полнота описания: {completeness}")
            if compliance_notes:
                additional_info_parts.append(f"Соответствие стандартам: {compliance_notes}")
            
            if isinstance(evidence_commentary, dict):
                if evidence_commentary.get('assessment'):
                    additional_info_parts.append(f"Оценка назначения: {evidence_commentary.get('assessment')}")
                if evidence_commentary.get('researchSummary'):
                    additional_info_parts.append(f"Резюме исследований: {evidence_commentary.get('researchSummary')}")
                if evidence_commentary.get('evidenceQuality'):
                    additional_info_parts.append(f"Качество доказательств: {evidence_commentary.get('evidenceQuality')}")
            
            additional_info = ' | '.join(additional_info_parts) if additional_info_parts else ''
            
            report_data.append({
                'Наименование в протоколе': r.get('original_name', 'N/A'),
                'МНН (англ)': r.get('inn', 'N/A'),
                'МНН (рус)': r.get('inn_russian', 'N/A'),
                'Дозировка': r.get('dosage', 'N/A'),
                'Способ введения': r.get('route', 'N/A'),
                'Кратность приема': r.get('frequency', 'N/A'),
                'Длительность': r.get('duration', 'N/A'),
                '💬 Резюме': summary_recommendation,
                '📋 Наличие в формулярах и одобрение регуляторов': status_summary,
                'УД (авторский)': r.get('author_ud', 'N/A'),
                '📊 УД (системный GRADE)': system_ud_display,
                '📝 Обоснование уровня доказательности': justification,
                '📄 Дополнительная информация': additional_info if additional_info else '✅ Нет дополнительной информации',
                '🔍 Руководство по проверке': evidence_commentary.get('evidenceSearchGuidanceText', '') if isinstance(evidence_commentary, dict) else '',
                '🔬 Исследования PubMed': pubmed_display,
                '🧪 Клинические исследования (ClinicalTrials.gov)': trials_display,
                '💭 Комментарии / Несоответствия': comments if comments else '✅ Нет замечаний'
            })

        # Create DataFrame
        df_main = pd.DataFrame(report_data)

        # Collect all PubMed articles for separate sheet
        pubmed_articles_data = []
        for r in all_verified_data:
            drug_name = r.get('original_name', 'N/A')
            inn_english = r.get('inn', 'N/A')
            pubmed_data = r.get('evidence_pubmed', {})
            articles = pubmed_data.get('articles', []) if isinstance(pubmed_data, dict) else []
            
            if articles:
                for article in articles:
                    pubmed_articles_data.append({
                        'Препарат': drug_name,
                        'МНН (англ)': inn_english,
                        'Название статьи': article.get('title', 'N/A'),
                        'Авторы': article.get('authors', 'N/A'),
                        'Журнал': article.get('journal', 'N/A'),
                        'Год': article.get('year', 'N/A'),
                        'Тип исследования': article.get('type', 'N/A'),
                        'PMID': article.get('pmid', 'N/A'),
                        'PMCID': article.get('pmcid', 'N/A'),  # PubMed Central ID (если доступен)
                        'DOI': article.get('doi', 'N/A'),  # Digital Object Identifier (если доступен)
                        'Ссылка': article.get('url', 'N/A')
                    })
        
        # Ensure the output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Create Excel writer for multiple sheets
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            # Main report sheet
            df_main.to_excel(writer, sheet_name='Анализ препаратов', index=False)
            
            # Protocol summary sheet if provided
            if protocol_summary:
                df_summary = pd.DataFrame({
                    'Executive Summary протокола': [protocol_summary]
                })
                df_summary.to_excel(writer, sheet_name='Резюме протокола', index=False)
            
            # PubMed articles sheet
            if pubmed_articles_data:
                df_pubmed = pd.DataFrame(pubmed_articles_data)
                df_pubmed.to_excel(writer, sheet_name='Исследования PubMed', index=False)
            
            # Auto-adjust column widths and apply conditional formatting
            from openpyxl.styles import PatternFill, Font, Alignment
            
            # Colors for GRADE
            fills = {
                '🟢': PatternFill(start_color="E6FFCC", end_color="E6FFCC", fill_type="solid"),  # Light Green for A
                '🟡': PatternFill(start_color="FFFFCC", end_color="FFFFCC", fill_type="solid"),  # Light Yellow for B
                '🟠': PatternFill(start_color="FFE5CC", end_color="FFE5CC", fill_type="solid"),  # Light Orange for C
                '🔴': PatternFill(start_color="FFCCCC", end_color="FFCCCC", fill_type="solid"),  # Light Red for D
                '⚪': PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")   # Grey for others
            }
            
            for sheet_name in writer.sheets:
                worksheet = writer.sheets[sheet_name]
                
                # Find headers
                headers = {}
                for cell in worksheet[1]:
                    headers[cell.value] = cell.column
                
                # Apply formatting
                grade_col_idx = headers.get('📊 УД (системный GRADE)')
                
                for row in worksheet.iter_rows(min_row=2):
                    for cell in row:
                        # Center alignment for all cells
                        cell.alignment = Alignment(vertical='center', wrap_text=True)
                        
                        # Apply conditional formatting for GRADE
                        if grade_col_idx and cell.column == grade_col_idx:
                            cell_val = str(cell.value)
                            for emoji, fill in fills.items():
                                if emoji in cell_val:
                                    cell.fill = fill
                                    break
                
                # Column widths
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter
                    for cell in column:
                        try:
                            if len(str(cell.value)) > max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass
                    adjusted_width = min(max_length + 2, 60) # Cap width at 60
                    worksheet.column_dimensions[column_letter].width = adjusted_width

        log.info(f"Enhanced report successfully saved to {output_path}")

    except Exception as e:
        log.error(f"Failed to generate the Excel report. Error: {e}")
        raise


def generate_word_report(all_verified_data: list[dict], output_path: Path, protocol_summary: str = None, template_path: Path = None):
    """
    Generates a Word (.docx) report from the verified data, using template if provided.
    
    Args:
        all_verified_data: The list of verification records with all analysis data.
        output_path: The Path object for the output .docx file.
        protocol_summary: Optional protocol executive summary.
        template_path: Optional path to template document (example.docx).
    """
    if not all_verified_data:
        log.warning("No data provided to generate Word report. Skipping file creation.")
        return
    
    log.info(f"Generating Word report for {len(all_verified_data)} records...")
    
    try:
        # Используем шаблон если указан
        if template_path and template_path.exists():
            log.info(f"Using template: {template_path}")
            doc = Document(str(template_path))
            
            # Очищаем содержимое шаблона, но сохраняем стили и настройки
            # Удаляем все параграфы кроме первого (заголовок)
            while len(doc.paragraphs) > 1:
                doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)
            
            # Удаляем все таблицы из шаблона (они будут заменены)
            for table in doc.tables[:]:
                table._element.getparent().remove(table._element)
            
            # Обновляем заголовок (если есть параграфы)
            if doc.paragraphs:
                # Сохраняем форматирование первого параграфа
                first_para = doc.paragraphs[0]
                first_para.text = 'АНАЛИЗ КЛИНИЧЕСКОГО ПРОТОКОЛА ПО ПРЕПАРАТАМ'
                first_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Убеждаемся что текст жирный
                if first_para.runs:
                    first_para.runs[0].font.bold = True
                else:
                    # Если нет runs, создаем новый
                    run = first_para.add_run('АНАЛИЗ КЛИНИЧЕСКОГО ПРОТОКОЛА ПО ПРЕПАРАТАМ')
                    run.font.bold = True
                    first_para.text = ''  # Очищаем старый текст
                    first_para.add_run('АНАЛИЗ КЛИНИЧЕСКОГО ПРОТОКОЛА ПО ПРЕПАРАТАМ').font.bold = True
        else:
            # Create new document
            doc = Document()
            
            # Set document margins (как в шаблоне)
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Title (в стиле шаблона)
            title_para = doc.add_paragraph('АНАЛИЗ КЛИНИЧЕСКОГО ПРОТОКОЛА ПО ПРЕПАРАТАМ')
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if title_para.runs:
                title_para.runs[0].font.bold = True
        
        # Protocol summary if available - parse markdown and format nicely
        if protocol_summary:
            doc.add_heading('Общее резюме протокола', 1)
            # Parse markdown-style summary and format properly
            lines = protocol_summary.split('\n')
            current_section = None
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # Check for headings
                if line.startswith('##'):
                    # Add spacing before new section
                    if current_section:
                        doc.add_paragraph()
                    current_section = line.replace('##', '').strip()
                    doc.add_heading(current_section, 2)
                elif line.startswith('#'):
                    doc.add_heading(line.replace('#', '').strip(), 1)
                elif line.startswith('- ') or line.startswith('• '):
                    # Bullet point - используем безопасный способ без стиля
                    para = doc.add_paragraph()
                    para.add_run('• ').bold = False
                    para.add_run(line[2:])
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
            doc.add_paragraph()  # Empty line before table
        
        # Добавляем пустую строку перед таблицей
        doc.add_paragraph()
        
        # Create table for drugs (в стиле шаблона - используем простой стиль таблицы)
        # Если используется шаблон, создаем таблицу в его стиле (5 колонок как в example.docx)
        use_template_style = template_path and template_path.exists()
        
        if use_template_style:
            # Таблица в стиле шаблона: 5 колонок с одинаковым заголовком
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'  # Простой стиль как в шаблоне
            
            # Header row (все колонки с одинаковым текстом как в шаблоне)
            header_cells = table.rows[0].cells
            header_text = 'Сведения о препаратах клинического протокола'
            for i in range(5):
                header_cells[i].text = header_text
                header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Убираем жирный шрифт (как в шаблоне)
                if header_cells[i].paragraphs[0].runs:
                    header_cells[i].paragraphs[0].runs[0].font.bold = False
            
            # Первая строка данных - общая информация
            info_row = table.add_row().cells
            info_text = f'В результате проведенной методологической оценки клинического протокола проанализировано {len(all_verified_data)} препаратов.'
            for i in range(5):
                info_row[i].text = info_text
                info_row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            # Стандартная таблица: 4 колонки
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Наименование препарата'
            header_cells[1].text = 'МНН'
            header_cells[2].text = 'УД (системный GRADE)'
            header_cells[3].text = 'Статус регистрации и одобрения'
            
            # Style header row
            for cell in header_cells:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add data rows
        for r in all_verified_data:
            row_cells = table.add_row().cells
            
            if use_template_style:
                # Формат как в шаблоне: каждая колонка содержит информацию о препарате
                drug_name = r.get('original_name', 'N/A')
                
                # Колонка 1: Наименование препарата
                row_cells[0].text = drug_name
                
                # Колонка 2: МНН
                inn_english = r.get('inn', 'N/A')
                inn_russian = r.get('inn_russian', '')
                if inn_russian:
                    inn_text = f"{inn_english}\n({inn_russian})"
                else:
                    inn_text = inn_english
                row_cells[1].text = inn_text
                
                # Колонка 3: УД (системный GRADE)
                system_ud = r.get('system_ud', 'D')
                if isinstance(system_ud, tuple):
                    system_ud_level = system_ud[0]
                else:
                    system_ud_level = str(system_ud)
                row_cells[2].text = f'УД: {system_ud_level}'
                
                # Колонка 4: Статус регистрации и одобрения
                status_parts = []
                if r.get('formulary_bnf'):
                    status_parts.append('BNF')
                if r.get('formulary_who_eml'):
                    status_parts.append('WHO EML')
                
                fda_data = r.get('approval_fda', {})
                is_fda_approved = fda_data.get('approved', False) if isinstance(fda_data, dict) else fda_data
                if is_fda_approved:
                    status_parts.append('FDA')
                
                if r.get('approval_ema') == 'Approved':
                    status_parts.append('EMA')
                
                kz_data = r.get('kz_register', {})
                kz_found = isinstance(kz_data, dict) and kz_data.get('found', False)
                if kz_found:
                    kz_details = ['Казахстан']
                    reg_number = kz_data.get('registration_number', '').strip()
                    reg_status = kz_data.get('registration_status', '').strip()
                    if reg_number:
                        kz_details.append(f"(№{reg_number}")
                        if reg_status:
                            kz_details.append(f", {reg_status})")
                        else:
                            kz_details.append(")")
                    elif reg_status:
                        kz_details.append(f"({reg_status})")
                    status_parts.append(''.join(kz_details))
                
                row_cells[3].text = ', '.join(status_parts) if status_parts else 'Не найдено'
                
                # Колонка 5: Дополнительная информация (дозировка, способ введения, исследования)
                info_parts = []
                if r.get('dosage'):
                    info_parts.append(f"Дозировка: {r.get('dosage')}")
                if r.get('route'):
                    info_parts.append(f"Способ: {r.get('route')}")
                
                pubmed_data = r.get('evidence_pubmed', {})
                if isinstance(pubmed_data, dict):
                    rct_count = pubmed_data.get('rct_count', 0)
                    meta_count = pubmed_data.get('meta_analysis_count', 0)
                    if rct_count > 0 or meta_count > 0:
                        research_info = []
                        if meta_count > 0:
                            research_info.append(f"Мета-анализы: {meta_count}")
                        if rct_count > 0:
                            research_info.append(f"РКИ: {rct_count}")
                        if research_info:
                            info_parts.append('; '.join(research_info))
                
                row_cells[4].text = '\n'.join(info_parts) if info_parts else 'Нет дополнительной информации'
                
                # Выравнивание всех ячеек как в шаблоне (JUSTIFY)
                for cell in row_cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                # Стандартный формат: 4 колонки
                # Drug name
                row_cells[0].text = r.get('original_name', 'N/A')
                
                # INN
                inn_english = r.get('inn', 'N/A')
                inn_russian = r.get('inn_russian', '')
                if inn_russian:
                    row_cells[1].text = f"{inn_english} ({inn_russian})"
                else:
                    row_cells[1].text = inn_english
                
                # System UD
                system_ud = r.get('system_ud', 'D')
                if isinstance(system_ud, tuple):
                    system_ud_level = system_ud[0]
                else:
                    system_ud_level = str(system_ud)
                row_cells[2].text = system_ud_level
                
                # Status summary
                fda_data = r.get('approval_fda', {})
                is_fda_approved = fda_data.get('approved', False) if isinstance(fda_data, dict) else fda_data
                
                kz_data = r.get('kz_register', {})
                kz_found = isinstance(kz_data, dict) and kz_data.get('found', False)
                
                status_parts = []
                if r.get('formulary_bnf'):
                    status_parts.append('BNF')
                if r.get('formulary_who_eml'):
                    status_parts.append('WHO EML')
                if is_fda_approved:
                    status_parts.append('FDA')
                if r.get('approval_ema') == 'Approved':
                    status_parts.append('EMA')
                if kz_found:
                    status_parts.append('Казахстан')
                
                row_cells[3].text = ', '.join(status_parts) if status_parts else 'Не найдено'
        
        # Добавляем пустую строку после таблицы
        doc.add_paragraph()
        
        # Add detailed information for each drug (без разрыва страницы, как в шаблоне)
        doc.add_paragraph('ДЕТАЛЬНАЯ ИНФОРМАЦИЯ ПО ПРЕПАРАТАМ')
        para = doc.paragraphs[-1]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if para.runs:
            para.runs[0].font.bold = True
        doc.add_paragraph()
        
        for idx, r in enumerate(all_verified_data, 1):
            # Drug heading
            drug_name = r.get('original_name', 'N/A')
            doc.add_heading(f'{idx}. {drug_name}', 2)
            
            # Basic info
            p = doc.add_paragraph()
            p.add_run('МНН (англ): ').bold = True
            p.add_run(r.get('inn', 'N/A'))
            
            inn_russian = r.get('inn_russian', '')
            if inn_russian:
                p = doc.add_paragraph()
                p.add_run('МНН (рус): ').bold = True
                p.add_run(inn_russian)
            
            if r.get('dosage'):
                p = doc.add_paragraph()
                p.add_run('Дозировка: ').bold = True
                p.add_run(r.get('dosage', 'N/A'))
            
            if r.get('route'):
                p = doc.add_paragraph()
                p.add_run('Способ введения: ').bold = True
                p.add_run(r.get('route', 'N/A'))
            
            # Evidence level
            system_ud = r.get('system_ud', 'D')
            if isinstance(system_ud, tuple):
                system_ud_level = system_ud[0]
            else:
                system_ud_level = str(system_ud)
            
            p = doc.add_paragraph()
            p.add_run('Уровень доказательности (GRADE): ').bold = True
            p.add_run(system_ud_level)
            
            justification = r.get('evidence_level_justification', '')
            if justification:
                p = doc.add_paragraph()
                p.add_run('Обоснование: ').bold = True
                p.add_run(justification)
            
            # Completeness and compliance (from HTML demo version)
            completeness = r.get('completeness', '')
            if completeness:
                p = doc.add_paragraph()
                p.add_run('Полнота описания: ').bold = True
                p.add_run(completeness)
            
            compliance_notes = r.get('compliance_notes', '')
            if compliance_notes:
                p = doc.add_paragraph()
                p.add_run('Замечания по соответствию стандартам: ').bold = True
                p.add_run(compliance_notes)
            
            # Evidence-based commentary (from HTML demo version)
            evidence_commentary = r.get('evidence_commentary', {})
            if isinstance(evidence_commentary, dict) and evidence_commentary:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('Комментарии по доказательной базе:').bold = True
                
                if evidence_commentary.get('assessment'):
                    p = doc.add_paragraph()
                    p.add_run('Оценка назначения: ').bold = True
                    p.add_run(evidence_commentary.get('assessment'))
                
                if evidence_commentary.get('generalNotes'):
                    p = doc.add_paragraph()
                    p.add_run('Общие замечания: ').bold = True
                    p.add_run(evidence_commentary.get('generalNotes'))
                
                if evidence_commentary.get('researchSummary'):
                    p = doc.add_paragraph()
                    p.add_run('Резюме исследований: ').bold = True
                    p.add_run(evidence_commentary.get('researchSummary'))
                
                if evidence_commentary.get('evidenceQuality'):
                    p = doc.add_paragraph()
                    p.add_run('Качество доказательств: ').bold = True
                    p.add_run(evidence_commentary.get('evidenceQuality'))
                
                if evidence_commentary.get('evidenceSearchGuidanceText'):
                    p = doc.add_paragraph()
                    p.add_run('Руководство по самостоятельной проверке доказательной базы:').bold = True
                    doc.add_paragraph(evidence_commentary.get('evidenceSearchGuidanceText'))
            
            # Summary recommendation
            summary_recommendation = r.get('summary_recommendation', '')
            if summary_recommendation:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('Резюме: ').bold = True
                p.add_run(summary_recommendation)
            
            # Regulatory status
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('Статус регистрации и одобрения:').bold = True
            
            status_info = []
            if r.get('formulary_bnf'):
                status_info.append('• Включен в BNF')
            if r.get('formulary_bnfc'):
                status_info.append('• Включен в BNF Children')
            if r.get('formulary_who_eml'):
                status_info.append('• Включен в WHO EML')
            
            fda_data = r.get('approval_fda', {})
            is_fda_approved = fda_data.get('approved', False) if isinstance(fda_data, dict) else fda_data
            if is_fda_approved:
                status_info.append('• Одобрен FDA')
            
            ema_status = r.get('approval_ema', '')
            if ema_status == 'Approved':
                status_info.append('• Одобрен EMA')
            
            kz_data = r.get('kz_register', {})
            if isinstance(kz_data, dict) and kz_data.get('found'):
                kz_status = "• Зарегистрирован в Казахстане"
                
                # Status from "Вид" column
                reg_status = kz_data.get('registration_status', '').strip()
                if reg_status:
                    kz_status += f" (Статус: {reg_status})"
                
                # Reregistration status
                rereg_status = kz_data.get('reregistration_status', '').strip()
                if rereg_status:
                    kz_status += f" (Перерегистрация: {rereg_status})"
                
                # Registration number
                reg_number = kz_data.get('registration_number', '').strip()
                if reg_number:
                    kz_status += f" (Номер: {reg_number})"
                
                # Dosage form from "Лек. форма" column
                dosage_form = kz_data.get('dosage_forms', '').strip()
                if dosage_form:
                    kz_status += f" (Форма: {dosage_form})"
                
                status_info.append(kz_status)
            
            if status_info:
                for info in status_info:
                    # Используем безопасный способ создания списка без стиля
                    para = doc.add_paragraph()
                    para.add_run('• ').bold = False
                    para.add_run(info)
            else:
                para = doc.add_paragraph()
                para.add_run('• ').bold = False
                para.add_run('Не найдено')
            
            # Research summary (counts only, no URLs)
            pubmed_data = r.get('evidence_pubmed', {})
            if isinstance(pubmed_data, dict):
                rct_count = pubmed_data.get('rct_count', 0)
                meta_count = pubmed_data.get('meta_analysis_count', 0)
                systematic_count = pubmed_data.get('systematic_review_count', 0)
                total_articles = pubmed_data.get('total_found', 0)
                
                if total_articles > 0 or rct_count > 0 or meta_count > 0:
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run('Исследования:').bold = True
                    research_info = []
                    if meta_count > 0:
                        research_info.append(f"Мета-анализы: {meta_count}")
                    if systematic_count > 0:
                        research_info.append(f"Систематические обзоры: {systematic_count}")
                    if rct_count > 0:
                        research_info.append(f"РКИ: {rct_count}")
                    if total_articles > 0:
                        research_info.append(f"Всего найдено исследований: {total_articles}")
                    if research_info:
                        doc.add_paragraph('; '.join(research_info))
            
            # Clinical trials (count only)
            trials = r.get('clinical_trials', [])
            if trials and len(trials) > 0:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('Клинические исследования: ').bold = True
                p.add_run(f"Найдено исследований: {len(trials)}")
            
            # Comments
            comments = r.get('comments', '')
            if comments:
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('Комментарии:').bold = True
                doc.add_paragraph(comments)
            
            # Add spacing between drugs
            doc.add_paragraph()
        
        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Save document
        doc.save(str(output_path))
        log.info(f"Word report successfully saved to {output_path}")
    
    except Exception as e:
        log.error(f"Failed to generate Word report. Error: {e}")
        raise
